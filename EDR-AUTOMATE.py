import os
import sys
import subprocess
import psutil
import hashlib
import time
import threading
import platform
from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler
from scapy.all import sniff
from docx import Document
from datetime import datetime

# ==========================
# Directory Structure Setup
# ==========================
def setup_directories():
    """Create necessary directories for logging and reporting."""
    base_dir = os.path.expanduser("~/EDR-System/")
    plexor_dir = os.path.join(base_dir, "PLEXOR")
    edr_dir = os.path.join(base_dir, "EDR")
    reporting_dir = os.path.join(base_dir, "Docx-Reporting")

    # Create directories if they don't exist
    os.makedirs(plexor_dir, exist_ok=True)
    os.makedirs(edr_dir, exist_ok=True)
    os.makedirs(reporting_dir, exist_ok=True)

    log_file = os.path.join(reporting_dir, 'monitoring_log.txt')
    # Create an empty log file for monitoring
    with open(log_file, 'w') as log:
        log.write(f"[INFO] Monitoring started at {datetime.now()}\n")

    print(f"[INFO] Folder structure set up at: {base_dir}")
    return plexor_dir, edr_dir, reporting_dir, log_file

# ==========================
# Dependency Installation
# ==========================
def install_dependencies():
    """Install required dependencies based on the user's platform."""
    try:
        print("[INFO] Installing necessary dependencies...")
        
        # Check if the platform is Windows, macOS, or Linux
        if platform.system() == "Windows":
            subprocess.check_call([sys.executable, "-m", "pip", "install", "psutil", "watchdog", "scapy", "python-docx"])
            print("[INFO] Dependencies installed for Windows.")
        elif platform.system() == "Darwin":  # macOS
            subprocess.check_call([sys.executable, "-m", "pip", "install", "psutil", "watchdog", "scapy", "python-docx"])
            print("[INFO] Dependencies installed for macOS.")
        elif platform.system() == "Linux":
            subprocess.check_call([sys.executable, "-m", "pip", "install", "psutil", "watchdog", "scapy", "python-docx"])
            print("[INFO] Dependencies installed for Linux.")
        else:
            print("[ERROR] Unsupported OS detected.")
            sys.exit(1)

    except subprocess.CalledProcessError as e:
        print(f"[ERROR] Failed to install dependencies: {str(e)}")
        sys.exit(1)

# ==========================
# System Process Monitoring
# ==========================
def monitor_system_processes(report_path):
    """Monitor system processes cross-platform and write results to the report."""
    with open(report_path, 'a') as report:
        report.write("[INFO] Monitoring system processes...\n")
        
    if platform.system() == "Windows":
        command = "tasklist"
    else:
        command = "ps aux"
    
    while True:
        result = subprocess.run(command, shell=True, stdout=subprocess.PIPE)
        with open(report_path, 'a') as report:
            report.write(result.stdout.decode())
        time.sleep(10)  # Monitor every 10 seconds

# ==========================
# File System Monitoring
# ==========================
class FileChangeHandler(FileSystemEventHandler):
    def __init__(self, report_path):
        self.report_path = report_path

    def on_modified(self, event):
        with open(self.report_path, 'a') as report:
            report.write(f'[ALERT] File modified: {event.src_path}\n')

    def on_created(self, event):
        with open(self.report_path, 'a') as report:
            report.write(f'[ALERT] File created: {event.src_path}\n')

    def on_deleted(self, event):
        with open(self.report_path, 'a') as report:
            report.write(f'[ALERT] File deleted: {event.src_path}\n')

def monitor_file_system(path_to_watch, report_path):
    """Monitor file changes in the specified directory and log to report."""
    event_handler = FileChangeHandler(report_path)
    observer = Observer()
    observer.schedule(event_handler, path=path_to_watch, recursive=True)
    observer.start()
    try:
        while True:
            time.sleep(1)
    except KeyboardInterrupt:
        observer.stop()
    observer.join()

# ==========================
# Network Traffic Monitoring
# ==========================
def packet_callback(packet, report_path):
    with open(report_path, 'a') as report:
        report.write(packet.summary() + '\n')

def monitor_network_traffic(report_path):
    """Monitor network traffic and write results to report."""
    sniff(prn=lambda packet: packet_callback(packet, report_path), count=100)  # Capture 100 packets

# ==========================
# Full Directory Malware Scanning
# ==========================
def scan_directory_for_malware(directory_to_scan, known_hashes, report_path):
    """Traverse the directory and scan each file for malware based on hash comparison."""
    for root, dirs, files in os.walk(directory_to_scan):
        for file in files:
            try:
                file_path = os.path.join(root, file)
                with open(file_path, 'rb') as f:
                    file_data = f.read()
                    file_hash = hashlib.md5(file_data).hexdigest()
                    with open(report_path, 'a') as report:
                        report.write(f"[INFO] Scanned file {file_path}, MD5: {file_hash}\n")
                    
                    if file_hash in known_hashes:
                        with open(report_path, 'a') as report:
                            report.write(f"[WARNING] Malware detected in {file_path} (local signature match)!\n")
            except (PermissionError, FileNotFoundError):
                with open(report_path, 'a') as report:
                    report.write(f"[ERROR] Could not access {file_path}\n")

# ==========================
# Forensic Reporting
# ==========================
def generate_forensic_report(report_dir, incident_details):
    """Generate a final forensic report using the logs."""
    report_file = os.path.join(report_dir, f'Forensic_Report_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx')
    print(f"[INFO] Generating forensic report at {report_file}...")
    doc = Document()
    doc.add_heading('Forensic Report', 0)
    doc.add_paragraph('Incident Summary:')
    doc.add_paragraph(incident_details)

    # Append log details from the monitoring process
    log_file = os.path.join(report_dir, 'monitoring_log.txt')
    with open(log_file, 'r') as log:
        doc.add_heading('Monitoring Details:', level=1)
        doc.add_paragraph(log.read())
    
    doc.save(report_file)
    print(f"[INFO] Forensic report saved at {report_file}.")

# ==========================
# Main Function: Entry Point
# ==========================
if __name__ == "__main__":
    # Install dependencies
    install_dependencies()

    # Set up folder structure
    plexor_dir, edr_dir, reporting_dir, log_file = setup_directories()

    # Example known malware hashes for signature-based detection
    known_malware_hashes = ['d41d8cd98f00b204e9800998ecf8427e']  # MD5 of empty file (replace with actual malware hashes)

    # Directory to monitor (Home directory or any other)
    path_to_watch = os.path.expanduser("~")

    # Incident Details (for final reporting)
    incident_details = "Suspicious activity detected during monitoring."

    print("[INFO] Starting EDR system...")

    # Start system process monitoring
    threading.Thread(target=monitor_system_processes, args=(log_file,), daemon=True).start()

    # Start file system monitoring
    threading.Thread(target=monitor_file_system, args=(path_to_watch, log_file), daemon=True).start()

    # Start network monitoring
    threading.Thread(target=monitor_network_traffic, args=(log_file,), daemon=True).start()

    # Start scanning the user's home directory for malware
    threading.Thread(target=scan_directory_for_malware, args=(os.path.expanduser("~"), known_malware_hashes, log_file), daemon=True).start()

    # Generate forensic report after monitoring
    time.sleep(60)  # Wait for some time to gather data before generating the report
    generate_forensic_report(reporting_dir, incident_details)
